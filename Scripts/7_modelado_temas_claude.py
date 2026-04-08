#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import os
import random
import re
from datetime import datetime
from pathlib import Path

import anthropic
from docx import Document
from docx.shared import Pt

try:
    from dotenv import load_dotenv

    env_file = Path(__file__).resolve().parent.parent / ".env.local"
    if env_file.exists():
        load_dotenv(str(env_file))
except ImportError:
    pass

from output_naming import build_output_dir, build_report_tag, ensure_tagged_name


REPO_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_INPUT_DIR = REPO_ROOT / "Datos"
DEFAULT_OUTPUT_DIR = REPO_ROOT / "Claude"
API_ENV_NAME = "CLAUDE_API_KEY"
DEFAULT_MODEL = "claude-opus-4-6"
DEFAULT_MAX_CORPUS_CHARS = 650000
DEFAULT_SAMPLE_SEED = 42
INPUT_FILENAMES = ("material_institucional.txt", "material_comentarios.txt")
MONTH_NAMES = {
    1: "ENERO",
    2: "FEBRERO",
    3: "MARZO",
    4: "ABRIL",
    5: "MAYO",
    6: "JUNIO",
    7: "JULIO",
    8: "AGOSTO",
    9: "SEPTIEMBRE",
    10: "OCTUBRE",
    11: "NOVIEMBRE",
    12: "DICIEMBRE",
}


def log_message(message: str) -> None:
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{timestamp}] {message}")


def valid_date(value: str) -> str:
    try:
        datetime.strptime(value, "%Y-%m-%d")
    except ValueError as exc:
        raise argparse.ArgumentTypeError(f"Fecha invalida '{value}', usa YYYY-MM-DD") from exc
    return value


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Combina los materiales de Datos, los envia a Claude y genera salidas semanales para Tampico"
    )
    parser.add_argument("--since", required=True, type=valid_date,
                        help="Fecha inicio YYYY-MM-DD (define la semana ISO)")
    parser.add_argument("--before", required=True, type=valid_date,
                        help="Fecha fin YYYY-MM-DD (compatibilidad con orquestador)")
    parser.add_argument("--input-dir", default=str(DEFAULT_INPUT_DIR),
                        help=f"Carpeta base de Datos (default: {DEFAULT_INPUT_DIR})")
    parser.add_argument("--output-dir", default=str(DEFAULT_OUTPUT_DIR),
                        help=f"Carpeta base de salida Claude (default: {DEFAULT_OUTPUT_DIR})")
    parser.add_argument("--model", default=DEFAULT_MODEL,
                        help=f"Modelo Claude a usar (default: {DEFAULT_MODEL})")
    parser.add_argument("--max-corpus-chars", type=int, default=DEFAULT_MAX_CORPUS_CHARS,
                        help=f"Maximo de caracteres a enviar (default: {DEFAULT_MAX_CORPUS_CHARS})")
    parser.add_argument("--sample-seed", type=int, default=DEFAULT_SAMPLE_SEED,
                        help=f"Semilla para muestreo aleatorio si el corpus excede el limite (default: {DEFAULT_SAMPLE_SEED})")
    parser.add_argument("--prepare-only", action="store_true",
                        help="Solo genera el corpus combinado y valida rutas, sin llamar a Claude")
    return parser.parse_args()


def weekly_input_dir(base_dir: Path, since: str) -> Path:
    return Path(base_dir) / build_report_tag(since, "Datos")


def weekly_output_dir(base_dir: Path, since: str) -> Path:
    return build_output_dir(base_dir, since, "Claude")


def read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            with open(path, "r", encoding=encoding, errors="ignore") as handle:
                return handle.read()
        except OSError:
            continue
    return ""


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as handle:
        handle.write(content)


def build_combined_corpus(datos_dir: Path, since: str) -> tuple[Path, str, list[str]]:
    datos_tag = build_report_tag(since, "Datos")
    corpus_name = f"{ensure_tagged_name('corpus_claude', datos_tag)}.txt"
    corpus_path = datos_dir / corpus_name

    sections: list[str] = []
    used_files: list[str] = []
    for filename in INPUT_FILENAMES:
        source_path = datos_dir / filename
        if not source_path.exists():
            log_message(f"⚠️ Archivo no encontrado, se omite: {source_path}")
            continue

        content = read_text(source_path).strip()
        if not content:
            log_message(f"⚠️ Archivo vacio, se omite: {source_path}")
            continue

        title = source_path.stem.replace("_", " ").upper()
        sections.append(f"=== {title} ===\n{content}")
        used_files.append(filename)

    if not sections:
        raise FileNotFoundError(
            f"No se encontraron materiales utilizables en {datos_dir}. Se esperaban: {', '.join(INPUT_FILENAMES)}"
        )

    combined_text = "\n\n".join(sections).strip() + "\n"
    write_text(corpus_path, combined_text)
    return corpus_path, combined_text, used_files


def sample_corpus(text: str, max_chars: int, seed: int) -> tuple[str, dict[str, int | bool | float]]:
    if len(text) <= max_chars:
        return text, {
            "sampled": False,
            "original_chars": len(text),
            "final_chars": len(text),
            "original_lines": len(text.splitlines()),
            "final_lines": len(text.splitlines()),
            "sample_ratio": 1.0,
        }

    lines = [line for line in text.splitlines() if line.strip()]
    if not lines:
        return text, {
            "sampled": False,
            "original_chars": len(text),
            "final_chars": len(text),
            "original_lines": 0,
            "final_lines": 0,
            "sample_ratio": 1.0,
        }

    ratio = max_chars / len(text)
    target_lines = max(1, int(len(lines) * ratio))
    rng = random.Random(seed)
    sampled_lines = rng.sample(lines, min(target_lines, len(lines)))
    sampled_text = "\n".join(sampled_lines)
    return sampled_text, {
        "sampled": True,
        "original_chars": len(text),
        "final_chars": len(sampled_text),
        "original_lines": len(lines),
        "final_lines": len(sampled_lines),
        "sample_ratio": round(len(sampled_lines) / len(lines), 4),
    }


def build_prompt(since: str) -> str:
    dt = datetime.strptime(since, "%Y-%m-%d")
    month_label = MONTH_NAMES[dt.month]
    year_label = dt.year
    return f"""
Analiza este corpus de conversación digital sobre TAMPICO, enfocándote específicamente en:
- Gestión del Gobierno Municipal de Tampico
- Acciones, decisiones y declaraciones de Mónica Villarreal Anaya, presidenta municipal de Tampico
- Políticas públicas, servicios urbanos e intervención municipal en Tampico (alumbrado, limpieza, seguridad, turismo, cultura, protección civil, infraestructura, movilidad y mantenimiento urbano)
- Percepción ciudadana sobre el gobierno municipal, sus funcionarios, programas y resultados
- Problemáticas, sucesos y debates propios de Tampico y su zona inmediata cuando afecten directamente al municipio

Genera un análisis temático con el siguiente formato EXACTO:

## ANÁLISIS TEMÁTICO DE LA CONVERSACIÓN DIGITAL SOBRE TAMPICO Y SU GOBIERNO MUNICIPAL - {month_label} DE {year_label}

### TEMAS PRINCIPALES (Narrativa descriptiva de los 8 temas más relevantes sobre Tampico y su gobierno municipal):

**1. [NOMBRE DEL TEMA EN MAYÚSCULAS]**
[Párrafo descriptivo de MÁXIMO 4 líneas (5 líneas máximo absoluto). Enfócate en acciones del gobierno municipal de Tampico, reacciones ciudadanas, debate público, cuestionamientos, respaldos o controversias específicas.]

**2. [NOMBRE DEL TEMA EN MAYÚSCULAS]**
[Párrafo descriptivo de MÁXIMO 4 líneas (5 líneas máximo absoluto).]

**3. [NOMBRE DEL TEMA EN MAYÚSCULAS]**
[Párrafo descriptivo de MÁXIMO 4 líneas (5 líneas máximo absoluto).]

**4. [NOMBRE DEL TEMA EN MAYÚSCULAS]**
[Párrafo descriptivo de MÁXIMO 4 líneas (5 líneas máximo absoluto).]

**5. [NOMBRE DEL TEMA EN MAYÚSCULAS]**
[Párrafo descriptivo de MÁXIMO 4 líneas (5 líneas máximo absoluto).]

**6. [NOMBRE DEL TEMA EN MAYÚSCULAS]**
[Párrafo descriptivo de MÁXIMO 4 líneas (5 líneas máximo absoluto).]

**7. [NOMBRE DEL TEMA EN MAYÚSCULAS]**
[Párrafo descriptivo de MÁXIMO 4 líneas (5 líneas máximo absoluto).]

**8. [NOMBRE DEL TEMA EN MAYÚSCULAS]**
[Párrafo descriptivo de MÁXIMO 4 líneas (5 líneas máximo absoluto).]

---

## ESTIMADO DE PESOS PORCENTUALES

**1. [TEMA EN MAYÚSCULAS] - XX.X%**
XXX menciones - [Descripción breve de 10-15 palabras en una sola línea]

**2. [TEMA EN MAYÚSCULAS] - XX.X%**
XXX menciones - [Descripción breve de 10-15 palabras en una sola línea]

**3. [TEMA EN MAYÚSCULAS] - XX.X%**
XXX menciones - [Descripción breve de 10-15 palabras en una sola línea]

**4. [TEMA EN MAYÚSCULAS] - XX.X%**
XXX menciones - [Descripción breve de 10-15 palabras en una sola línea]

**5. [TEMA EN MAYÚSCULAS] - XX.X%**
XXX menciones - [Descripción breve de 10-15 palabras en una sola línea]

**6. [TEMA EN MAYÚSCULAS] - XX.X%**
XXX menciones - [Descripción breve de 10-15 palabras en una sola línea]

**7. [TEMA EN MAYÚSCULAS] - XX.X%**
XXX menciones - [Descripción breve de 10-15 palabras en una sola línea]

**8. [TEMA EN MAYÚSCULAS] - XX.X%**
XXX menciones - [Descripción breve de 10-15 palabras en una sola línea]

**RESTRICCIONES ESTRICTAS DE FORMATO:**
- Cada narrativa de tema: MÁXIMO 4 líneas de texto, absolutamente no más de 5 líneas
- Cada descripción breve: MÁXIMO 10-15 palabras, UNA SOLA LÍNEA
- Nombres de temas en MAYÚSCULAS
- EXACTAMENTE 8 temas en ambas secciones, ni más ni menos
- Los 8 porcentajes deben sumar 100%
- NO incluyas categoría "Otros temas" o similar
- Formato sin tabla, solo lista numerada con formato:
  **N. TEMA - XX.X%**
  XXX menciones - Descripción

**CRITERIOS DE CONTENIDO:**
- Identifica EXACTAMENTE 8 temas principales que cubran todo el corpus
- PRIORIZA temas relacionados con:
  * Gestión y decisiones del Gobierno Municipal de Tampico
  * Acciones, declaraciones y programas de Mónica Villarreal Anaya y su administración
  * Servicios públicos municipales (alumbrado, limpieza, bacheo, mantenimiento urbano, protección civil, turismo, cultura, seguridad y atención ciudadana)
  * Infraestructura y desarrollo urbano de Tampico
  * Percepción ciudadana sobre el desempeño del Ayuntamiento de Tampico
- Si aparecen actores estatales o federales, inclúyelos solo cuando impacten directamente la conversación sobre Tampico
- Usa lenguaje técnico pero accesible para tomadores de decisiones
- Categoriza por temática específica de Tampico y su gobierno municipal, no por sentimiento general
- Las menciones deben ser números reales estimados del corpus
- Ignora conversaciones ajenas a Tampico o que no tengan impacto claro en la agenda municipal
""".strip()


def generar_analisis_claude(api_key: str, model: str, prompt: str, corpus_text: str) -> tuple[str, dict[str, int | str]]:
    client = anthropic.Anthropic(api_key=api_key)
    message_content = f"{prompt}\n\n=== CORPUS PARA ANÁLISIS ===\n\n{corpus_text}"

    log_message("🚀 Enviando corpus a Claude API...")
    response = client.messages.create(
        model=model,
        max_tokens=4000,
        messages=[{"role": "user", "content": message_content}],
    )

    analysis_text = "\n".join(
        block.text for block in response.content if getattr(block, "type", "") == "text"
    ).strip()
    usage = {
        "model": model,
        "input_tokens": int(getattr(response.usage, "input_tokens", 0)),
        "output_tokens": int(getattr(response.usage, "output_tokens", 0)),
    }
    return analysis_text, usage


def create_word_document(analysis_text: str, output_path: Path) -> None:
    doc = Document()

    for raw_line in analysis_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=1)
            continue

        if line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=2)
            continue

        if re.match(r"\*\*\d+\.", line):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line.replace("**", ""))
            run.bold = True
            run.font.size = Pt(11)
            continue

        if line.startswith("---"):
            doc.add_paragraph()
            continue

        doc.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def save_metadata(path: Path, metadata: dict[str, object]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as handle:
        json.dump(metadata, handle, ensure_ascii=False, indent=2)
        handle.write("\n")


def main() -> None:
    args = parse_args()

    input_dir = Path(args.input_dir)
    output_dir = Path(args.output_dir)
    datos_dir = weekly_input_dir(input_dir, args.since)
    claude_dir = weekly_output_dir(output_dir, args.since)
    claude_tag = build_report_tag(args.since, "Claude")

    log_message("🤖 ANÁLISIS TEMÁTICO CON CLAUDE PARA TAMPICO")
    log_message(f"Semana de datos: {datos_dir}")
    log_message(f"Salida Claude: {claude_dir}")

    if not datos_dir.exists():
        raise SystemExit(f"No existe la carpeta semanal de datos: {datos_dir}")

    corpus_path, corpus_text, used_files = build_combined_corpus(datos_dir, args.since)
    log_message(f"📄 Corpus combinado generado: {corpus_path.name}")
    log_message(f"📚 Archivos integrados: {', '.join(used_files)}")

    sampled_corpus, sampling_stats = sample_corpus(corpus_text, args.max_corpus_chars, args.sample_seed)
    if sampling_stats["sampled"]:
        log_message(
            f"⚠️ Corpus excede el limite; se aplico muestreo aleatorio a {sampling_stats['final_lines']} lineas "
            f"de {sampling_stats['original_lines']}"
        )
    else:
        log_message(f"✅ Corpus completo listo para envio: {sampling_stats['final_chars']:,} caracteres")

    if args.prepare_only:
        log_message("🧪 Modo prepare-only: no se llamo a Claude")
        return

    api_key = os.getenv(API_ENV_NAME, "").strip()
    if not api_key:
        raise SystemExit(
            f"No se encontro {API_ENV_NAME}. Define la variable en .env.local o en el entorno antes de ejecutar este script."
        )

    prompt = build_prompt(args.since)
    analysis_text, usage = generar_analisis_claude(api_key, args.model, prompt, sampled_corpus)
    if not analysis_text:
        raise SystemExit("Claude no devolvio texto de analisis")

    md_path = claude_dir / f"{ensure_tagged_name('analisis_tematico_claude', claude_tag)}.md"
    docx_path = claude_dir / f"{ensure_tagged_name('analisis_tematico_claude', claude_tag)}.docx"
    meta_path = claude_dir / f"{ensure_tagged_name('metadata_claude', claude_tag)}.json"

    write_text(md_path, analysis_text.strip() + "\n")
    create_word_document(analysis_text, docx_path)
    save_metadata(
        meta_path,
        {
            "since": args.since,
            "before": args.before,
            "input_dir": str(datos_dir),
            "output_dir": str(claude_dir),
            "combined_corpus": str(corpus_path),
            "input_files": used_files,
            "sampling": sampling_stats,
            "usage": usage,
        },
    )

    log_message(f"✅ Analisis guardado en: {md_path}")
    log_message(f"✅ Documento Word guardado en: {docx_path}")
    log_message(f"✅ Metadata guardada en: {meta_path}")


if __name__ == "__main__":
    main()